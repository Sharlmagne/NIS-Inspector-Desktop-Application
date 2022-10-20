from report import Report
from docx import Document


class Itinerary:
    def __init__(self, report: Report):
        self.name = None
        self.grade = None
        self.period_start = None
        self.period_end = None
        self.area = None
        self.parish_office = None
        self.objectives = []
        self.report = report

    def update_report(self):
        # Update the reports objectives
        self.report.objectives = self.objectives

    def update_itinerary_document(self, filename):
        # Load document
        doc = Document('files/blank_itinerary.docx')

        # Update document heading
        doc.paragraphs[2].clear().add_run("  PARISH OFFICE: ")
        doc.paragraphs[2].add_run(self.parish_office).bold = True
        doc.paragraphs[4].clear().add_run("NAME OF OFFICER: ")
        doc.paragraphs[4].add_run(self.name).bold = True
        doc.paragraphs[4].add_run("			  GRADE: ")
        doc.paragraphs[4].add_run(self.grade).bold = True
        doc.paragraphs[5].clear().add_run("PERIOD: ")
        doc.paragraphs[5].add_run(f"{self.period_start} - {self.period_end}").bold = True
        doc.paragraphs[5].add_run("			  AREA: ")
        doc.paragraphs[5].add_run(self.area).bold = True

        # Get the table object
        table = doc.tables[0]

        # Load the table with all the objectives from the itinerary object
        for row, i in zip(table.rows[2:len(self.objectives) + 2], range(len(self.objectives))):
            row.cells[0].text = self.objectives[i]['date']
            row.cells[1].text = self.objectives[i]['business_name']
            row.cells[2].text = self.objectives[i]['business_address']
            row.cells[3].text = self.objectives[i]['visit']

        try:
            doc.save(filename)
        except PermissionError:
            return False
        else:
            return True
