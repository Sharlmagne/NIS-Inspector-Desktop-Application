from docx import Document
from docx.shared import Pt

results = ["Effective", "Ineffective"]


def font_edit(text, font_size):
    text.paragraphs[0].runs[0].font.size = Pt(font_size)
    text.paragraphs[0].runs[0].font.name = 'Times New Roman'
    text.paragraphs[0].runs[0].font.bold = True


def load_report_table(table, obj):
    # Load each row of the report table with the objectives from the object
    for row, i in zip(table.rows[1:len(obj) + 1], range(len(obj))):
        row.cells[0].text = obj[i]['date']
        row.cells[1].text = obj[i]['business_address']
        row.cells[2].text = obj[i]['business_name']
        row.cells[3].text = obj[i]['ref_no']
        row.cells[4].text = obj[i]['visit']
        row.cells[6].text = obj[i]['result'].capitalize()

        font_edit(row.cells[0], 12)
        font_edit(row.cells[1], 9)
        font_edit(row.cells[2], 9)
        font_edit(row.cells[3], 12)
        font_edit(row.cells[4], 12)
        font_edit(row.cells[6], 12)


class Report:
    def __init__(self):
        self.officer = None
        self.week_ended = None
        self.area = None
        self.files_carried_fwd = None
        self.files_received = None
        self.total = None
        self.files_cleared = None
        self.files_on_hand = None
        self.objectives = []
        self.total_effective = None
        self.total_ineffective = None
        self.ineffective = {
            "BEN.": 0, "R.I.": 0, "COMP.": 0, "REG.": 0, "EDUC.": 0, "C/L": 0, "TOTAL": 0
        }
        self.effective = {
            "BEN.": 0, "R.I.": 0, "COMP.": 0, "REG.": 0, "EDUC.": 0, "C/L": 0, "TOTAL": 0
        }
        self.results_total = {
                "BEN.": 0, "R.I.": 0, "COMP.": 0, "REG.": 0, "EDUC.": 0,  "C/L": 0,
            }

    def objective_results(self):
        # Count the number of effective and ineffective objectives
        self.reset_results()
        for obj in self.objectives:
            if obj['result'] == "effective" and obj['visit'] in self.results_total.keys():
                self.effective[obj["visit"]] += 1
            elif obj['result'] == "ineffective" and obj['visit'] in self.results_total.keys():
                self.ineffective[obj["visit"]] += 1

        # Calculate the total results for the objectives
        ineffective_total = []
        effective_total = []
        for value in self.ineffective.values():
            ineffective_total.append(value)
        for value in self.effective.values():
            effective_total.append(value)
        self.ineffective["TOTAL"] = sum(ineffective_total)
        self.effective["TOTAL"] = sum(effective_total)
        self.total_ineffective = sum(ineffective_total)
        self.total_effective = sum(effective_total)
        self.result_count()

    def result_count(self):
        # Sum the total of effective and ineffective objectives per visit
        for obj in self.results_total:
            self.results_total[obj] = self.effective[obj] + self.ineffective[obj]

    def reset_results(self):
        # Reset the results count
        for key in self.effective:
            self.effective[key] = 0
            self.ineffective[key] = 0

    def load_report_heading(self, table):
        # Set the heading for the report word document
        table.clear().add_run("NAME OF OFFICER ")
        table.add_run(f"{self.officer}" + " " * 69).bold = True
        table.add_run("WEEK ENDED ")
        table.add_run(f"{self.week_ended}" + " " * 61).bold = True
        table.add_run("AREA ")
        table.add_run(f"{self.area}").bold = True

    def load_results_table(self, table):
        # Load the report table with the objectives
        for i, key in zip(range(1, 7), self.effective.keys()):
            table.rows[1].cells[i].text = str(self.effective[key])
            table.rows[2].cells[i].text = str(self.ineffective[key])
            table.rows[3].cells[i].text = str(self.results_total[key])
        table.rows[1].cells[7].text = str(self.total_effective)
        table.rows[2].cells[7].text = str(self.total_ineffective)

    def update_report_document(self, filename):
        report = Document('files/blank_report.docx')

        table = report.tables
        results_table = table[3]

        # Update the heading of the report document
        self.load_report_heading(report.paragraphs[0])
        self.load_report_heading(report.paragraphs[4])
        self.load_report_heading(report.paragraphs[8])

        number_of_objectives = len(self.objectives)

        # Check the number of objectives and update the three tables accordingly
        if number_of_objectives < 16:
            load_report_table(table[0], self.objectives[:])
        if 16 < number_of_objectives < 33:
            load_report_table(table[0], self.objectives[:16])
            load_report_table(table[1], self.objectives[16:])
        else:
            load_report_table(table[0], self.objectives[:16])
            load_report_table(table[1], self.objectives[16:])
            load_report_table(table[2], self.objectives[32:])
        self.load_results_table(results_table)

        # Update the files section of the report
        report.paragraphs[16].add_run("Number of files carried forward from previous week" + " " * 53 + f"{self.files_carried_fwd}")
        report.paragraphs[17].add_run("Number of files received during the week" + " " * 72 + f"{self.files_received}")
        report.paragraphs[19].clear().add_run("TOTAL" + " " * 55 + f"{self.total}")
        report.paragraphs[21].add_run("Number of files cleared during the week" + " " * 74 + f"{self.files_cleared}")
        report.paragraphs[22].add_run("Number of files on hand at the end of the week" + " " * 61 + f"{self.files_on_hand}")

        try:
            report.save(filename)
        except PermissionError:
            return False
        else:
            return True
